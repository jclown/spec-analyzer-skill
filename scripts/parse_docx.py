#!/usr/bin/env python3
"""
Word Document Parser
Parse .docx format specification documents
"""

import sys
import re
import json
import subprocess
from pathlib import Path
from typing import Dict, List, Optional
from dataclasses import dataclass, asdict
from enum import Enum


class Severity(Enum):
    MANDATORY = "mandatory"
    RECOMMENDED = "recommended"
    REFERENCE = "reference"
    UNKNOWN = "unknown"


@dataclass
class SpecRule:
    """Specification rule data structure"""
    id: str
    name: str
    category: str
    content: str
    severity: str
    positive_example: Optional[str] = None
    negative_example: Optional[str] = None


def install_dependency(package: str) -> bool:
    """
    Automatically install missing dependency package
    Returns True if installation succeeds, False otherwise
    """
    try:
        print(f"Installing dependency: {package}...")
        result = subprocess.run(
            [sys.executable, "-m", "pip", "install", package],
            capture_output=True,
            text=True,
            timeout=60
        )
        if result.returncode == 0:
            print(f"Successfully installed {package}")
            return True
        else:
            print(f"Failed to install {package}: {result.stderr}", file=sys.stderr)
            return False
    except subprocess.TimeoutExpired:
        print(f"Installation timed out for {package}", file=sys.stderr)
        return False
    except Exception as e:
        print(f"Error installing {package}: {e}", file=sys.stderr)
        return False


def parse_docx(file_path: str, auto_install: bool = True) -> Dict:
    """
    Parse Word document
    If auto_install is True, will automatically install python-docx if missing
    """
    try:
        from docx import Document
    except ImportError:
        if auto_install:
            if install_dependency("python-docx"):
                # Try importing again after installation
                try:
                    from docx import Document
                except ImportError:
                    return {
                        'error': 'Failed to import python-docx even after installation',
                        'fallback': 'Please convert document to Markdown format'
                    }
            else:
                return {
                    'error': 'Failed to install python-docx automatically',
                    'fallback': 'Please run: pip install python-docx, or convert document to Markdown format'
                }
        else:
            return {
                'error': 'python-docx not installed. Please run: pip install python-docx',
                'fallback': 'Please convert document to Markdown format'
            }

    doc = Document(file_path)

    # Extract all paragraph text
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Detect paragraph style to determine hierarchy level
            style = para.style.name if para.style else ''
            paragraphs.append({
                'text': text,
                'style': style,
                'level': _get_heading_level(style)
            })

    # Extract table content
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        if table_data:
            tables.append(table_data)

    # Parse into rules
    rules = _extract_rules(paragraphs)
    categories = _extract_categories(paragraphs)

    return {
        'file': file_path,
        'total_rules': len(rules),
        'categories': categories,
        'rules': [asdict(r) for r in rules],
        'tables': tables
    }


def _get_heading_level(style: str) -> int:
    """Get heading level from style name"""
    if not style:
        return 0
    # Check English heading styles
    if 'Heading 1' in style:
        return 1
    elif 'Heading 2' in style:
        return 2
    elif 'Heading 3' in style:
        return 3
    # Check Chinese heading styles (common variants)
    if '标题 1' in style or '标题1' in style:
        return 1
    elif '标题 2' in style or '标题2' in style:
        return 2
    elif '标题 3' in style or '标题3' in style:
        return 3
    return 0


def _extract_categories(paragraphs: List[Dict]) -> Dict[str, List[str]]:
    """Extract categories from paragraphs"""
    categories = {}
    current_category = "General"
    rule_ids = []

    for p in paragraphs:
        text = p['text']
        level = p['level']

        # Level 1 headings as categories
        if level == 1:
            if rule_ids:
                categories[current_category] = rule_ids
            current_category = text
            rule_ids = []
        elif level == 2:
            # Chinese numeral prefixes like "一、", "二、", "三、"
            numeral_match = re.match(r'^[一二三四五六七八九十]+[、．.]\s*', text)
            if numeral_match:
                if rule_ids:
                    categories[current_category] = rule_ids
                current_category = re.sub(r'^[一二三四五六七八九十]+[、．.]?\s*', '', text)
                rule_ids = []

    if rule_ids:
        categories[current_category] = rule_ids

    return categories


def _extract_rules(paragraphs: List[Dict]) -> List[SpecRule]:
    """Extract rules from paragraphs"""
    rules = []
    current_category = "General"
    rule_counter = 0

    i = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        text = p['text']

        # Update category
        if p['level'] in [1, 2]:
            current_category = re.sub(r'^[一二三四五六七八九十]+[、．.]?\s*', '', text)

        # Detect severity level
        severity = _detect_severity(text)
        if severity != Severity.UNKNOWN.value:
            rule_counter += 1

            # Extract rule ID
            rule_id_match = re.match(r'(\d+\.?\d*)', text)
            rule_id = rule_id_match.group(1) if rule_id_match else f"{rule_counter}"

            # Extract rule name and content
            content = re.sub(r'^(\d+\.?\d*\s*)?', '', text)
            content = re.sub(r'^【[^】]+】', '', content).strip()

            # Find example code
            positive = None
            negative = None

            # Look ahead for examples
            j = i + 1
            while j < len(paragraphs) and j < i + 10:
                next_text = paragraphs[j]['text']
                if '正例' in next_text or '正确' in next_text or '示例' in next_text or 'positive' in next_text.lower():
                    # Find code block
                    if j + 1 < len(paragraphs):
                        positive = paragraphs[j + 1]['text']
                elif '反例' in next_text or '错误' in next_text or '禁止' in next_text or 'negative' in next_text.lower():
                    if j + 1 < len(paragraphs):
                        negative = paragraphs[j + 1]['text']
                elif _detect_severity(next_text) != Severity.UNKNOWN.value:
                    break  # Encountered new rule
                j += 1

            rule = SpecRule(
                id=rule_id,
                name=content[:50] if len(content) > 50 else content,
                category=current_category,
                content=content,
                severity=severity,
                positive_example=positive,
                negative_example=negative
            )
            rules.append(rule)

        i += 1

    return rules


def _detect_severity(text: str) -> str:
    """Detect severity level from text"""
    # Chinese markers
    if re.search(r'【强制】', text):
        return Severity.MANDATORY.value
    if re.search(r'【推荐】', text):
        return Severity.RECOMMENDED.value
    if re.search(r'【参考】', text):
        return Severity.REFERENCE.value
    # English markers
    if re.search(r'\*\*MUST\*\*', text, re.IGNORECASE):
        return Severity.MANDATORY.value
    if re.search(r'\*\*SHOULD\*\*', text, re.IGNORECASE):
        return Severity.RECOMMENDED.value
    if re.search(r'\*\*MAY\*\*', text, re.IGNORECASE):
        return Severity.REFERENCE.value
    # Keyword detection (Chinese)
    for kw in ['禁止', '不得', '不要', '严禁']:
        if kw in text:
            return Severity.MANDATORY.value
    for kw in ['必须', '应当', '需要']:
        if kw in text:
            return Severity.MANDATORY.value
    for kw in ['建议', '推荐', '最好']:
        if kw in text:
            return Severity.RECOMMENDED.value
    # Keyword detection (English)
    if re.search(r'\bMUST\b', text, re.IGNORECASE):
        return Severity.MANDATORY.value
    if re.search(r'\bSHOULD\b', text, re.IGNORECASE):
        return Severity.RECOMMENDED.value
    if re.search(r'\bMAY\b', text, re.IGNORECASE):
        return Severity.REFERENCE.value
    return Severity.UNKNOWN.value


def main():
    if len(sys.argv) < 2:
        print("Usage: python parse_docx.py <docx_file> [output_file]")
        print("Options:")
        print("  --no-auto-install    Disable automatic dependency installation")
        sys.exit(1)

    file_path = sys.argv[1]
    output_path = None
    auto_install = True

    # Parse arguments
    for arg in sys.argv[2:]:
        if arg == '--no-auto-install':
            auto_install = False
        elif not arg.startswith('--'):
            output_path = arg

    result = parse_docx(file_path, auto_install)

    if 'error' in result:
        print(f"Error: {result['error']}", file=sys.stderr)
        if 'fallback' in result:
            print(f"Suggestion: {result['fallback']}", file=sys.stderr)
        sys.exit(1)

    output = json.dumps(result, ensure_ascii=False, indent=2)

    if output_path:
        Path(output_path).write_text(output, encoding='utf-8')
        print(f"Parsing complete: {result['total_rules']} rules extracted")
    else:
        print(output)


if __name__ == '__main__':
    main()